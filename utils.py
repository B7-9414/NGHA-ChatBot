import os
import pdfplumber
from docx import Document


import pdfplumber

def read_pdf(file_path):
    full_text = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            # Extract paragraph-like text
            text = page.extract_text()
            if text:
                full_text.append(text.strip())

            # Extract tables
            tables = page.extract_tables()
            for table in tables:
                for row in table:
                    clean_row = [cell.strip() for cell in row if cell and cell.strip()]
                    if clean_row:
                        full_text.append(" | ".join(clean_row))  # Markdown-style row

    return "\n".join(full_text)


def read_docx(file_path):
    from docx import Document
    doc = Document(file_path)
    full_text = []

    # Extract paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            full_text.append(text)

    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                full_text.append(" | ".join(cells))

    result = "\n".join(full_text)
    print("\n\n📄 DOCX Content Preview:\n", result[:1000])  # Only show first 1000 characters
    return result
