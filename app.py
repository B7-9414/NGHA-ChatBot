# app.py
import os
import time
import base64
import hashlib
import pdfplumber
from docx import Document as DocxDocument

import streamlit as st
from dotenv import load_dotenv

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LCDocument
from langchain_openai import ChatOpenAI
from langchain.chains import create_history_aware_retriever, create_retrieval_chain
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain_core.messages import AIMessage, HumanMessage

from langchain_chroma import Chroma
from langchain_ollama import OllamaEmbeddings
import streamlit.components.v1 as components
import tiktoken  # ✅ for token counting

# =========================
# Speak button (Web Speech API)
# =========================
def speak_button(text: str, key: str):
    """
    Renders a 🔉/🔇 toggle button that reads/stops text.
    key must be unique per message.
    Auto-detects Arabic to pick an Arabic TTS voice if available.
    """
    safe = (text or "").replace("\\", "\\\\").replace('"', '\\"').replace("'", "\\'")
    components.html(
        f"""
        <div>
          <button id="speak-btn-{key}"
                  onclick="toggleSpeak_{key}()"
                  style="background:none;border:none;cursor:pointer;font-size:20px">🔉</button>
        </div>
        <script>
          (function(){{  /* IIFE start */
            let utterance_{key} = null;
            let isSpeaking_{key} = false;
            const btn = document.getElementById("speak-btn-{key}");
            // crude Arabic detection
            const isArabic = /[\\u0600-\\u06FF]/.test("{safe}");
            function pickVoice(langStart) {{
              const voices = window.speechSynthesis.getVoices() || [];
              // try an exact match first, then startsWith
              let v = voices.find(v => v.lang === langStart) ||
                      voices.find(v => (v.lang || "").toLowerCase().startsWith(langStart.toLowerCase()));
              return v || null;
            }}
            function buildUtterance() {{
              const u = new SpeechSynthesisUtterance("{safe}");
              // Prefer Arabic voice if Arabic text detected
              if (isArabic) {{
                const ar = pickVoice("ar");
                if (ar) u.voice = ar;
                u.lang = "ar-SA";
              }} else {{
                const en = pickVoice("en");
                if (en) u.voice = en;
                u.lang = "en-US";
              }}
              u.rate = 1.0;   // speed
              u.pitch = 1.0;  // tone
              u.onend = () => {{
                isSpeaking_{key} = false;
                if (btn) btn.textContent = "🔉";
              }};
              return u;
            }}
            function start() {{
              // stop anything currently speaking (global)
              window.speechSynthesis.cancel();
              utterance_{key} = buildUtterance();
              window.speechSynthesis.speak(utterance_{key});
              isSpeaking_{key} = true;
              btn.textContent = "🔇";
            }}
            function stop() {{
              window.speechSynthesis.cancel();
              isSpeaking_{key} = false;
              btn.textContent = "🔉";
            }}
            window.toggleSpeak_{key} = function() {{
              if (!isSpeaking_{key}) start(); else stop();
            }};
            // Some browsers need this to populate voices list
            if (speechSynthesis && speechSynthesis.onvoiceschanged !== undefined) {{
              speechSynthesis.onvoiceschanged = function(){{}};
            }}
          }})(); /* IIFE end */
        </script>
        """,
        height=40
    )

# =========================
# Config
# =========================
UPLOAD_DIR = "data/uploads"
PERSIST_DIR = "data/chroma"
EMBED_MODEL = "llama3"          # Ollama embeddings
RETRIEVAL_K = 5
CHUNK_SIZE = 1200
CHUNK_OVERLAP = 200

# Use a constant for the LLM model so we can reuse for token counting
LLM_MODEL = "gpt-3.5-turbo"

load_dotenv()

# =========================
# UI & Styling
# =========================
st.set_page_config(
    page_title="NGHA | TAIF AI Assistant",
    page_icon="logo_image_en.png",
    layout="centered"
)

st.markdown("""
    <style>
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header {visibility: hidden;}
    .suggestion-container {
        display: flex;
        flex-direction: column;
        align-items: center;
        justify-content: center;
        margin-top: 30px;
        margin-bottom: 30px;
        animation: fadeIn 0.5s ease-in-out;
    }
    @keyframes fadeIn {
        from { opacity: 0; transform: translateY(20px); }
        to { opacity: 1; transform: translateY(0); }
    }
    </style>
""", unsafe_allow_html=True)

def get_base64_logo(path: str) -> str:
    with open(path, "rb") as f:
        return base64.b64encode(f.read()).decode()

# Load images
top_right_b64 = get_base64_logo("1.jpeg")        # Top-right (main page)
bottom_left_main = get_base64_logo("2.jpeg")     # Bottom-left (main page)
bottom_left_b64 = get_base64_logo("2.jpeg")      # Bottom-left (injected from sidebar)

# CSS and fixed-position images on the main page
st.markdown(
    f"""
    <style>
    /* ===== SIDEBAR CONTEXT (needed even if not positioning inside) ===== */
    [data-testid="stSidebar"] {{
        position: relative !important;
        overflow: visible !important;
    }}

    /* ===== FIXED IMAGE POSITIONS ON MAIN PAGE ===== */
    .bg-top-right {{
        position: fixed;
        top: 0;
        right: 0;
        width: 220px;
        opacity: 0.18;
        pointer-events: none;
        z-index: 1;
    }}

    .bg-bottom-left-main {{
        position: fixed;
        bottom: 110px;
        left: 10px;
        width: 180px;
        opacity: 0.18;
        pointer-events: none;
        z-index: 1;
    }}

    @media (max-width: 900px) {{
        .bg-bottom-left-main {{
            bottom: 140px;
            width: 150px;
        }}
    }}

    /* ===== SIDEBAR IMAGE (still positioned fixed) ===== */
    [data-testid="stSidebar"] .sb-bottom-left {{
        position: fixed;
        bottom: 110px;
        left: 10px;
        width: 180px;
        opacity: 0.18;
        pointer-events: none;
        z-index: 1;
    }}
    </style>

    <!-- Page-based images -->
    <img class="bg-top-right" src="data:image/jpeg;base64,{top_right_b64}">
    <img class="bg-bottom-left-main" src="data:image/jpeg;base64,{bottom_left_main}">
    """,
    unsafe_allow_html=True
)

# Inject the fixed-position image from inside the sidebar
with st.sidebar:
    st.markdown(
        f'<img class="sb-bottom-left" src="data:image/jpeg;base64,{bottom_left_b64}">',
        unsafe_allow_html=True
    )

# Title
st.title("NGHA | TAIF AI Assistant 🤖 🏥")

# =========================
# Running totals (tokens & cost) header (under title)
# =========================
for k, v in {
    "total_input_tokens": 0,
    "total_output_tokens": 0,
    "total_cost": 0.0
}.items():
    if k not in st.session_state:
        st.session_state[k] = v

header_stats = st.empty()

def render_header_stats():
    total_tokens = st.session_state["total_input_tokens"] + st.session_state["total_output_tokens"]
    header_stats.markdown(
        f"<div style='font-size:12px;color:#666;margin-top:-10px;'>"
        f"🔹 Tokens: {total_tokens} | 💰 Cost: ${st.session_state['total_cost']:.5f}"
        f"</div>",
        unsafe_allow_html=True
    )

render_header_stats()

# =========================
# Helpers: File IO & Split
# =========================
def read_pdf(file_path: str) -> str:
    parts = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                parts.append(text.strip())
            tables = page.extract_tables()
            for table in tables or []:
                for row in table:
                    row = [c.strip() for c in row if c and str(c).strip()]
                    if row:
                        parts.append(" | ".join(row))
    return "\n".join(parts)

def read_docx(file_path: str) -> str:
    doc = DocxDocument(file_path)
    parts = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)

def file_text(file_path: str) -> str:
    if file_path.lower().endswith(".pdf"):
        return read_pdf(file_path)
    elif file_path.lower().endswith(".docx"):
        return read_docx(file_path)
    else:
        raise ValueError("Unsupported file type (only .pdf, .docx).")

def content_sha256(path: str) -> str:
    """Hash file content to detect changes reliably."""
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()

def split_to_documents(text: str, meta: dict) -> list[LCDocument]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        separators=["\n\n", "\n", " ", ""],
    )
    chunks = splitter.split_text(text)
    docs = []
    for i, ch in enumerate(chunks):
        d = LCDocument(page_content=ch, metadata={**meta, "chunk_index": i})
        docs.append(d)
    return docs

# =========================
# Vector DB: Sync
# =========================
def get_vectordb():
    embeddings = OllamaEmbeddings(model=EMBED_MODEL)
    vectordb = Chroma(persist_directory=PERSIST_DIR, embedding_function=embeddings)
    return vectordb

def current_upload_files() -> dict[str, str]:
    """Return {file_id(hash): path} for .pdf/.docx in UPLOAD_DIR."""
    os.makedirs(UPLOAD_DIR, exist_ok=True)
    files = [os.path.join(UPLOAD_DIR, f) for f in os.listdir(UPLOAD_DIR)
             if f.lower().endswith((".pdf", ".docx"))]
    mapping = {}
    for p in files:
        try:
            fid = content_sha256(p)  # content-based ID
            mapping[fid] = p
        except Exception:
            continue
    return mapping

def vectordb_file_ids(vdb: Chroma) -> set[str]:
    fids = set()
    offset = 0
    page = 5000
    while True:
        results = vdb.get(include=["metadatas"], limit=page, offset=offset)
        metas = results.get("metadatas", []) or []
        if not metas:
            break
        for m in metas:
            fid = m.get("file_id")
            if isinstance(fid, str) and fid:
                fids.add(fid)
        if len(metas) < page:
            break
        offset += page
    return fids

def delete_by_file_id(vdb: Chroma, file_id: str):
    vdb.delete(where={"file_id": {"$eq": file_id}})

def add_file_to_vectordb(vdb: Chroma, file_path: str, file_id: str):
    text = file_text(file_path)
    if not text.strip():
        return
    meta = {
        "source": file_path,
        "file_name": os.path.basename(file_path),
        "file_id": file_id,
    }
    docs = split_to_documents(text, meta)
    vdb.add_documents(docs)

def sync_vectordb() -> Chroma:
    """
    - Adds new/changed files in UPLOAD_DIR
    - Removes vectors for files deleted from UPLOAD_DIR
    """
    vdb = get_vectordb()

    current_map = current_upload_files()             # {file_id: path}
    current_ids = set(current_map.keys())

    existing_ids = vectordb_file_ids(vdb)            # from DB

    # Delete missing
    for orphan_id in (existing_ids - current_ids):
        delete_by_file_id(vdb, orphan_id)

    # Add new
    for new_id in (current_ids - existing_ids):
        add_file_to_vectordb(vdb, current_map[new_id], new_id)

    try:
        vdb.persist()
    except Exception:
        pass

    return vdb

# =========================
# Build/Load Chains
# =========================
@st.cache_resource(show_spinner="🔄 Syncing documents ...")
def load_retriever():
    vdb = sync_vectordb()
    return vdb.as_retriever(search_kwargs={"k": RETRIEVAL_K})

llm = ChatOpenAI(model=LLM_MODEL, temperature=0.1)

search_prompt = ChatPromptTemplate.from_messages([
    MessagesPlaceholder("chat_history"),
    ("user", "{input}"),
    ("user", "Given the above conversation, generate a search query to look up in order to get information relevant to the conversation.")
])

answer_prompt = ChatPromptTemplate.from_messages([
    ("system", "Answer the user's questions based on the context below. Always be concise and relevant.\n\n{context}"),
    MessagesPlaceholder("chat_history"),
    ("user", "{input}")
])

retriever = load_retriever()
retriever_chain = create_history_aware_retriever(llm=llm, retriever=retriever, prompt=search_prompt)
stuff_chain = create_stuff_documents_chain(llm=llm, prompt=answer_prompt)
rag_chain = create_retrieval_chain(retriever_chain, stuff_chain)

# =========================
# Sidebar
# =========================
with st.sidebar:
    st.image("logo_image_en.png", width=100)
    st.markdown("### 📂 Uploaded Documents")
    os.makedirs(UPLOAD_DIR, exist_ok=True)
    files = [f for f in os.listdir(UPLOAD_DIR) if f.lower().endswith((".pdf", ".docx"))]
    if files:
        for f in sorted(files):
            st.markdown(f"* {f}")
    else:
        st.info("No documents found.")

    col1, col2 = st.columns(2)
    with col1:
        if st.button("🧹 New Chat"):
            # ✅ Let Streamlit's automatic rerun handle the refresh
            st.session_state.chat_history = []
            st.session_state.show_suggestions = True
            # reset running totals
            st.session_state.total_input_tokens = 0
            st.session_state.total_output_tokens = 0
            st.session_state.total_cost = 0.0
            # clear handles (don't .empty() here to avoid a blank frame)
            st.session_state.ai_extras = []
            # clear any pending suggestion
            st.session_state.pop("suggested_query", None)
            st.success("Started a new chat.")
            # ❌ Do NOT call st.rerun() here

    with col2:
        if st.button("🔄 Sync now"):
            # clear cache_resource and resync
            load_retriever.clear()
            _ = load_retriever()
            st.success("✅ Vector store synced successfully. Refreshing to reflect changes.")
            time.sleep(2)  # ⏱ gives user time to read the message
            st.rerun()  # refresh file list immediately

# =========================
# Session State
# =========================
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []
if "show_suggestions" not in st.session_state:
    st.session_state.show_suggestions = True
if "ai_extras" not in st.session_state:
    st.session_state.ai_extras = []  # holds placeholders for last AI speaker+tokens

def _hide_suggestions_on_submit():
    st.session_state.show_suggestions = False

# =========================
# Suggestions (fast hide)
# =========================
sugg_box = st.empty()  # placeholder for the whole suggestions area

def render_suggestions():
    with sugg_box.container():
        st.markdown("""
        <div class="suggestion-container">
            <h4 style='color:#cccccc;'>🧠 Suggested Questions:</h4>
        </div>
        """, unsafe_allow_html=True)

        suggestions = [
            "What is the extension of Thomas Bright?",
            "List all emails ending with @gmail.com",
            "ما هو رقم الشارة لـ Victoria Ashley؟",
            "اعرض جميع الإيميلات التي تنتهي بـ @gmail.com"
        ]
        rows = [suggestions[:2], suggestions[2:]]
        for row in rows:
            cols = st.columns(len(row), gap="large")
            for col, question in zip(cols, row):
                with col:
                    if st.button(question, use_container_width=True, key=question):
                        # 1) stash the query in session
                        st.session_state.suggested_query = question
                        st.session_state.show_suggestions = False
                        # 2) hide immediately on the client (no waiting)
                        sugg_box.empty()
                        # 3) rerun so the rest of the app picks it up
                        st.rerun()

# Only render suggestions when appropriate
if st.session_state.get("show_suggestions", True) \
   and len(st.session_state.get("chat_history", [])) == 0 \
   and "suggested_query" not in st.session_state:
    render_suggestions()

# =========================
# Chat History Rendering
# =========================
for msg in st.session_state.chat_history:
    role = "AI" if isinstance(msg, AIMessage) else "Human"
    avatar = "logo_image_en.png" if role == "AI" else None
    with st.chat_message(role, avatar=avatar):
        st.write(msg.content)

# =========================
# Token counter helper
# =========================
def count_tokens(text: str, model: str = LLM_MODEL) -> int:
    try:
        enc = tiktoken.encoding_for_model(model)
    except Exception:
        enc = tiktoken.get_encoding("cl100k_base")
    return len(enc.encode(text or ""))

# =========================
# Chat Input & RAG Call
# =========================
user_query = st.chat_input(
    "Ask about any of the uploaded PDFs/DOCX ...",
    key="chatbox",
    on_submit=_hide_suggestions_on_submit
)

# --- Handle suggestion click: rerun with suggestion ---
if "suggested_query" in st.session_state and not user_query:
    try:
        sugg_box.empty()   # hide suggestions instantly
    except Exception:
        pass
    user_query = st.session_state.pop("suggested_query")
    st.session_state.show_suggestions = False

if user_query:
    # hide suggestions immediately (typed or suggestion)
    try:
        sugg_box.empty()
    except Exception:
        pass
    st.session_state.show_suggestions = False

    # 🔻 Instantly remove previous speaker & token rows (no lag)
    for ph in st.session_state.get("ai_extras", []):
        try:
            ph.empty()
        except Exception:
            pass
    st.session_state.ai_extras = []

    # Show the user's message
    st.session_state.chat_history.append(HumanMessage(content=user_query))
    with st.chat_message("Human"):
        st.write(user_query)

    input_tokens = count_tokens(user_query, model=LLM_MODEL)

    # Get AI answer
    with st.chat_message("AI", avatar="logo_image_en.png"):
        placeholder = st.empty()
        with st.spinner("Thinking..."):
            response = rag_chain.invoke({
                "input": user_query,
                "chat_history": st.session_state.chat_history
            })
            final_answer = response["answer"] if isinstance(response, dict) else str(response)

        # typing effect
        typed = ""
        for word in final_answer.split():
            typed += word + " "
            placeholder.markdown(typed + "▌")
            time.sleep(0.03)
        final_text = typed.strip()
        placeholder.markdown(final_text)

        # === Speaker + Token line in ONE container (fast clear next time) ===
        extras_box = st.empty()
        with extras_box.container():
            # 🔉 Speak button
            msg_index = sum(1 for m in st.session_state.chat_history if isinstance(m, AIMessage))
            speak_button(final_text, key=f"{msg_index}")

            # 💰 Cost summary
            output_tokens = count_tokens(final_text, model=LLM_MODEL)
            price_per_1k_input = 0.0005
            price_per_1k_output = 0.0015
            input_cost = (input_tokens / 1000.0) * price_per_1k_input
            output_cost = (output_tokens / 1000.0) * price_per_1k_output
            st.session_state.total_input_tokens += input_tokens
            st.session_state.total_output_tokens += output_tokens
            st.session_state.total_cost += (input_cost + output_cost)
            render_header_stats()

            st.markdown(
                f"<div style='font-size:10px;color:#888;'>"
                f"🔹 Tokens: in {input_tokens} + out {output_tokens} = {input_tokens + output_tokens} "
                f"| 💰 Cost: ${(input_cost + output_cost):.5f}"
                f"</div>",
                unsafe_allow_html=True
            )

        # Keep a handle so we can clear fast next time
        st.session_state.ai_extras.append(extras_box)

        # Save assistant message
        st.session_state.chat_history.append(AIMessage(content=final_text))
